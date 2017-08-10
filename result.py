#
#GNU General Public Licence v3.0
#Copyright (c) 2017 Vivek Bhagat
#
#
#This program is free software: you can redistribute it and/or modify
#it under the terms of the GNU General Public License as published by
#the Free Software Foundation, either version 3 of the License, or
#(at your option) any later version.
#
#This program is distributed in the hope that it will be useful,
#but WITHOUT ANY WARRANTY; without even the implied warranty of
#MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
#GNU General Public License for more details.
#
#You should have received a copy of the GNU General Public License
#along with this program.  If not, see <http://www.gnu.org/licenses/>.


from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
#from docx.enum.style import WD_STYLE_TYPE
from openpyxl import load_workbook
import random


##Creating an object for Document()
result = Document()


##Loading excel document to read data co-scholastic grades
wb = load_workbook('co-scholistic grade.xlsx')
#wb = load_workbook('testfile.xlsx')
ws = wb['Sheet1']
#Loading data to a 2 dimensional list
lstgrade = []
for row in ws.iter_rows(min_row=2, max_col=22, max_row=100):
    valuerow = []
    for cell in row:
        valuerow.append(cell.value)
    lstgrade.append(valuerow)


##Loading excel document to read data discriptive indicator(M)
wb = load_workbook('discriptive indicator (M).xlsx')
ws = wb['Sheet1']
#Loading data to a 2 dimensional list
lstindicatorM = []
for row in ws.iter_rows(min_row=1, max_col=1, max_row=46):
    #valuerow = []
    for cell in row:
        lstindicatorM.append(cell.value)


##Loading excel document to read data discriptive indicator (F)
wb = load_workbook('discriptive indicator (F).xlsx')
ws = wb['Sheet1']
#Loading data to a 2 dimensional list
lstindicatorF = []
for row in ws.iter_rows(min_row=1, max_col=1, max_row=46):
    #valuerow = []
    for cell in row:
        lstindicatorF.append(cell.value)


##Adding a picture school letter head
#result.add_picture('name.png', width=Inches(1.25))


##Modifying the page layout
sections = result.sections[0]
sections.top_margin = Inches(1.5)
sections.bottom_margin = Inches(0)
sections.left_margin = Inches(0.5)
sections.right_margin = Inches(0.5)
#sections.header_distance = Pt(0)
#sections.footer_distance = Pt(0)


##Loop to iterate all students
for i in range(1, 99):

    ##Storing record for each student iteration by iteration
    srecord = lstgrade[i]
    lstindicator = []

    ##Checking weather the student is male or female
    if srecord[21] == 'M':
        lstindicator = lstindicatorM
    else:
        lstindicator = lstindicatorF

    ##Heading
    heading = result.add_paragraph()
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.add_run('Co-Scholastic Grade Certificate Class X 2017').bold = True
    #head1=result.add_heading('Co-Scholastic Grade Certificate Class X 2017',3).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


    ##Body setting
    paragraph_format = result.styles['Normal'].paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.line_spacing = Pt(10)
    #paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style = result.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)


    ##Personal Info (Body)
    para1 = result.add_paragraph('\n')
    para1.add_run('\nName of Student: ')
    #Cleaning the data for extra spaces and make all the text upper case
    info = " ".join(srecord[3].split())
    para1.add_run(info.upper())
    para1.add_run('\nAdmission No: ')
    info = " ".join(srecord[1].split())
    para1.add_run(info.upper())
    para1.add_run('\t\t\t\tSection: ')
    info = " ".join(srecord[2].split())
    para1.add_run(info.upper())
    para1.add_run('\nRoll No: ')
    info = " ".join(srecord[0].split())
    para1.add_run(info.upper())
    para1.add_run("\nMother's name: ")
    info = " ".join(srecord[4].split())
    para1.add_run(info.upper())
    para1.add_run("\nFather's name: ")
    info = " ".join(srecord[5].split())
    para1.add_run(info.upper())


    ##2(A) Life Skills:
    para2a = result.add_paragraph()
    para2a.add_run('2(A) Life Skills:')
    #Creating table with 4 rows and 4 columns
    table2a = result.add_table(rows=4, cols=4)
    table2a.style = 'TableGrid'
    #Setting the column width
    table2a.columns[0].width = Inches(1.3)
    table2a.columns[1].width = Inches(4.8)
    table2a.columns[2].width = Inches(0.6)
    table2a.columns[3].width = Inches(1)
    table2a.cell(0,0).text = 'Life Skills'
    table2a.cell(0,1).text = 'Descriptive Indicator'
    table2a.cell(0,2).text = 'Grade'
    table2a.cell(0,3).text = 'Grade Point'
    table2a.cell(1,0).text = 'Thinking Skills'
    table2a.cell(2,0).text = 'Social Skills'
    table2a.cell(3,0).text = 'Emotional Skills'
    #Thinking Skills data
    table2a.cell(1,3).text = str(srecord[7]).split()
    if srecord[7] == 5:
        drecord = [0,1]
        table2a.cell(1,2).text = 'A'
        table2a.cell(1,1).text = str(" ".join(lstindicator[random.choice(drecord)].split()))
    else:
        drecord = [2,3]
        table2a.cell(1,2).text = 'B'
        table2a.cell(1,1).text = str(" ".join(lstindicator[random.choice(drecord)].split()))
    #Social Skills data
    table2a.cell(2,3).text = str(srecord[8]).split()
    if srecord[8] == 5:
        drecord = [4,5]
        table2a.cell(2,2).text = 'A'
        table2a.cell(2,1).text = str(" ".join(lstindicator[random.choice(drecord)].split()))
    else:
        drecord = [6,7]
        table2a.cell(2,2).text = 'B'
        table2a.cell(2,1).text = str(" ".join(lstindicator[random.choice(drecord)].split()))
    #Emotional Skills data
    table2a.cell(3,3).text = str(srecord[9]).split()
    if srecord[9] == 5:
        drecord = [8,9]
        table2a.cell(3,2).text = 'A'
        table2a.cell(3,1).text = str(" ".join(lstindicator[random.choice(drecord)].split()))
    else:
        drecord = [10]
        table2a.cell(3,2).text = 'B'
        table2a.cell(3,1).text = str(" ".join(lstindicator[random.choice(drecord)].split()))


    ##2(B) Work Education:
    para2b = result.add_paragraph()
    para2b.add_run('\n2(B)')
    #Creating table with 2 rows and 4 columns
    table2b = result.add_table(rows=2, cols=4)
    table2b.style = 'TableGrid'
    #Setting the column width
    table2b.columns[0].width = Inches(1.3)
    table2b.columns[1].width = Inches(4.8)
    table2b.columns[2].width = Inches(0.6)
    table2b.columns[3].width = Inches(1)
    table2b.cell(0,0).text = 'Work Education'
    table2b.cell(0,1).text = 'Descriptive Indicator'
    table2b.cell(0,2).text = 'Grade'
    table2b.cell(0,3).text = 'Grade Point'
    table2b.cell(1,0).text = 'Work Education'
    #Work Education data
    table2b.cell(1,3).text = str(srecord[10]).split()
    if srecord[10] == 5:
        drecord = [11,12]
        table2b.cell(1,2).text = 'A'
        table2b.cell(1,1).text = str(" ".join(lstindicator[random.choice(drecord)].split()))
    else:
        drecord = [13,14]
        table2b.cell(1,2).text = 'B'
        table2b.cell(1,1).text = str(" ".join(lstindicator[random.choice(drecord)].split()))


    ##2(C) Visual and Performing Arts:
    para2c = result.add_paragraph()
    para2c.add_run('\n2(C)')
    #Creating table with 2 rows and 4 columns
    table2c = result.add_table(rows=2, cols=4)
    table2c.style = 'TableGrid'
    #Setting the column width
    table2c.columns[0].width = Inches(1.3)
    table2c.columns[1].width = Inches(4.8)
    table2c.columns[2].width = Inches(0.6)
    table2c.columns[3].width = Inches(1)
    table2c.cell(0,1).text = 'Descriptive Indicator'
    table2c.cell(0,2).text = 'Grade'
    table2c.cell(0,3).text = 'Grade Point'
    table2c.cell(1,0).text = 'Visual and Performing Arts'
    #Visual and Performing Arts data
    table2c.cell(1,3).text = str(srecord[11]).split()
    if srecord[11] == 5:
        drecord = [15,16]
        table2c.cell(1,2).text = 'A'
        table2c.cell(1,1).text = str(" ".join(lstindicator[random.choice(drecord)].split()))
    else:
        drecord = [17]
        table2c.cell(1,2).text = 'B'
        table2c.cell(1,1).text = str(" ".join(lstindicator[random.choice(drecord)].split()))


    ##2(D) Attitudes and Values:
    para2d = result.add_paragraph()
    para2d.add_run('\n2(D) Attitudes and Values:')
    #Creating table with 5 rows and 4 columns
    table2d = result.add_table(rows=5, cols=4)
    table2d.style = 'TableGrid'
    #Setting the column width
    table2d.columns[0].width = Inches(1.3)
    table2d.columns[1].width = Inches(4.8)
    table2d.columns[2].width = Inches(0.6)
    table2d.columns[3].width = Inches(1)
    table2d.cell(0,0).text = 'Attitude towards'
    table2d.cell(0,1).text = 'Descriptive Indicator'
    table2d.cell(0,2).text = 'Grade'
    table2d.cell(0,3).text = 'Grade Point'
    table2d.cell(1,0).text = 'Teachers'
    table2d.cell(2,0).text = 'Schoolmates'
    table2d.cell(3,0).text = 'School Programmes & Environment'
    table2d.cell(4,0).text = 'Value Systems'
    #Teachers data
    table2d.cell(1,3).text = str(srecord[12]).split()
    if srecord[12] == 5:
        drecord = [18,19]
        table2d.cell(1,2).text = 'A'
        table2d.cell(1,1).text = str(" ".join(lstindicator[random.choice(drecord)].split()))
    else:
        drecord = [20]
        table2d.cell(1,2).text = 'B'
        table2d.cell(1,1).text = str(" ".join(lstindicator[random.choice(drecord)].split()))
    #Schoolmates data
    table2d.cell(2,3).text = str(srecord[13]).split()
    if srecord[13] == 5:
        drecord = [21,22]
        table2d.cell(2,2).text = 'A'
        table2d.cell(2,1).text = str(" ".join(lstindicator[random.choice(drecord)].split()))
    else:
        drecord = [23]
        table2d.cell(2,2).text = 'B'
        table2d.cell(2,1).text = str(" ".join(lstindicator[random.choice(drecord)].split()))
    #School Programmes & Environment data
    table2d.cell(3,3).text = str(srecord[14]).split()
    if srecord[14] == 5:
        drecord = [24,25]
        table2d.cell(3,2).text = 'A'
        table2d.cell(3,1).text = str(" ".join(lstindicator[random.choice(drecord)].split()))
    else:
        drecord = [26]
        table2d.cell(3,2).text = 'B'
        table2d.cell(3,1).text = str(" ".join(lstindicator[random.choice(drecord)].split()))
    #Value Systems
    table2d.cell(4,3).text = str(srecord[15]).split()
    if srecord[15] == 5:
        drecord = [27,28]
        table2d.cell(4,2).text = 'A'
        table2d.cell(4,1).text = str(" ".join(lstindicator[random.choice(drecord)].split()))
    else:
        drecord = [29]
        table2d.cell(4,2).text = 'B'
        table2d.cell(4,1).text = str(" ".join(lstindicator[random.choice(drecord)].split()))


    ##3(A) Co-Curricular Activities:
    para3a = result.add_paragraph()
    para3a.add_run('\n3(A) Co-Curricular Activities:')
    #Creating table with 3 rows and 4 columns
    table3a = result.add_table(rows=3, cols=4)
    table3a.style = 'TableGrid'
    #Setting the column width
    table3a.columns[0].width = Inches(1.3)
    table3a.columns[1].width = Inches(4.8)
    table3a.columns[2].width = Inches(0.6)
    table3a.columns[3].width = Inches(1)
    table3a.cell(0,0).text = 'Activity'
    table3a.cell(0,1).text = 'Descriptive Indicator'
    table3a.cell(0,2).text = 'Grade'
    table3a.cell(0,3).text = 'Grade Point'
    table3a.cell(1,0).text = 'Literacy & Creative Skills'
    table3a.cell(2,0).text = 'Scientific Skills'
    #Literacy and Creative Skills data
    table3a.cell(1,3).text = str(srecord[16]).split()
    if srecord[16] == 5:
        drecord = [30,31]
        table3a.cell(1,2).text = 'A'
        table3a.cell(1,1).text = str(" ".join(lstindicator[random.choice(drecord)].split()))
    else:
        drecord = [32,33]
        table3a.cell(1,2).text = 'B'
        table3a.cell(1,1).text = str(" ".join(lstindicator[random.choice(drecord)].split()))
    #Scientific Skills data
    table3a.cell(2,3).text = str(srecord[17]).split()
    if srecord[17] == 5:
        drecord = [34,35]
        table3a.cell(2,2).text = 'A'
        table3a.cell(2,1).text = str(" ".join(lstindicator[random.choice(drecord)].split()))
    else:
        drecord = [36]
        table3a.cell(2,2).text = 'B'
        table3a.cell(2,1).text = str(" ".join(lstindicator[random.choice(drecord)].split()))


    ##3(B) Physical and Health Education:
    para3b = result.add_paragraph()
    para3b.add_run('\n3(B)Physical and Health Education:')
    #Creating table with 3 rows and 4 columns
    table3b = result.add_table(rows=3, cols=4)
    table3b.style = 'TableGrid'
    #Setting the column width
    table3b.columns[0].width = Inches(1.3)
    table3b.columns[1].width = Inches(4.8)
    table3b.columns[2].width = Inches(0.6)
    table3b.columns[3].width = Inches(1)
    table3b.cell(0,0).text = 'Activity'
    table3b.cell(0,1).text = 'Descriptive Indicator'
    table3b.cell(0,2).text = 'Grade'
    table3b.cell(0,3).text = 'Grade Point'
    table3b.cell(1,0).text = 'Sports'
    table3b.cell(2,0).text = 'Gardening & Shramdaan'
    #Sports data
    table3b.cell(1,3).text = str(srecord[18]).split()
    if srecord[18] == 5:
        drecord = [37,38,39]
        table3b.cell(1,2).text = 'A'
        table3b.cell(1,1).text = str(" ".join(lstindicator[random.choice(drecord)].split()))
    else:
        drecord = [40,41]
        table3b.cell(1,2).text = 'B'
        table3b.cell(1,1).text = str(" ".join(lstindicator[random.choice(drecord)].split()))
    #Gardening and Shramdaan data
    table3b.cell(2,3).text = str(srecord[19]).split()
    if srecord[19] == 5:
        drecord = [42,43]
        table3b.cell(2,2).text = 'A'
        table3b.cell(2,1).text = str(" ".join(lstindicator[random.choice(drecord)].split()))
    else:
        drecord = [44,45]
        table3b.cell(2,2).text = 'B'
        table3b.cell(2,1).text = str(" ".join(lstindicator[random.choice(drecord)].split()))


    ##Total
    Para4 = result.add_paragraph('\n\t\t\t\t\t\t\t\t\t\t\t\tTotal = ')
    Para4.add_run(str(srecord[20]).split())


    ##Adding page break
    result.add_page_break()


##Saving Document
result.save('result.docx')
